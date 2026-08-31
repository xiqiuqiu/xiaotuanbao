#!/usr/bin/env python3
"""从仓库源码生成软著用「前30页+后30页」源代码 docx，并统计源程序量。"""

from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[3]
OUT = Path(__file__).resolve().parents[1] / "output" / "小团宝旅行社运营管理系统V1.0-源代码.docx"
SOFTWARE_HEADER = "小团宝旅行社运营管理系统 V1.0"

# 每页行数（软著要求约 ≥50）
LINES_PER_PAGE = 50
PAGES_EACH_END = 30
CHUNK_LINES = LINES_PER_PAGE * PAGES_EACH_END  # 1500

SOURCE_ROOTS = [
    ROOT / "apps" / "api" / "src",
    ROOT / "apps" / "api" / "prisma",
    ROOT / "packages" / "shared" / "src",
    ROOT / "apps" / "web" / "src",
]

EXTENSIONS = {".ts", ".tsx", ".js", ".jsx", ".css", ".prisma", ".sql"}

# 启动入口优先排在最前
BOOTSTRAP = ROOT / "apps" / "api" / "src" / "main.ts"

SKIP_DIR_NAMES = {
    "node_modules",
    "dist",
    "build",
    "coverage",
    "__snapshots__",
    ".git",
}


def set_run_font(run, *, size_pt: float = 9) -> None:
    run.font.size = Pt(size_pt)
    run.font.name = "Courier New"
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), "Courier New")
    rFonts.set(qn("w:hAnsi"), "Courier New")
    rFonts.set(qn("w:eastAsia"), "宋体")


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, size_pt=9)


def setup_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    # 窄边距以便每页塞满约 50 行
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(1.8)

    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.clear()
    tab_stops = hp.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(17.0), WD_TAB_ALIGNMENT.RIGHT)
    run = hp.add_run(f"{SOFTWARE_HEADER}    代码")
    set_run_font(run, size_pt=9)
    hp.add_run("\t")
    add_page_number(hp)

    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p.clear()


def iter_source_files() -> list[Path]:
    files: list[Path] = []
    seen: set[Path] = set()

    def add(path: Path) -> None:
        path = path.resolve()
        if path in seen or not path.is_file():
            return
        if path.suffix not in EXTENSIONS:
            return
        seen.add(path)
        files.append(path)

    if BOOTSTRAP.is_file():
        add(BOOTSTRAP)

    collected: list[Path] = []
    for root in SOURCE_ROOTS:
        if not root.is_dir():
            continue
        for dirpath, dirnames, filenames in os.walk(root):
            dirnames[:] = [d for d in dirnames if d not in SKIP_DIR_NAMES and not d.startswith(".")]
            for name in sorted(filenames):
                p = Path(dirpath) / name
                if p.suffix in EXTENSIONS:
                    # 软著材料可含测试文件；若只需业务代码可在此过滤 *.test.*
                    collected.append(p)

    collected.sort(key=lambda p: str(p.relative_to(ROOT)).replace("\\", "/"))
    for p in collected:
        add(p)
    return files


def read_nonblank_lines(path: Path) -> list[str]:
    try:
        text = path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        text = path.read_text(encoding="utf-8", errors="ignore")
    out: list[str] = []
    for line in text.splitlines():
        # 去掉空行；保留缩进代码
        if line.strip() == "":
            continue
        # 控制单行过长，避免版式爆炸
        if len(line) > 120:
            line = line[:120]
        out.append(line.rstrip())
    return out


def collect_all_lines(files: list[Path]) -> tuple[list[str], int, int]:
    """返回 (全部非空行, 文件数, 含空行的物理行数近似统计用非空行数)."""
    all_lines: list[str] = []
    for path in files:
        rel = path.relative_to(ROOT).as_posix()
        body = read_nonblank_lines(path)
        if not body:
            continue
        # 文件分隔注释，便于阅读且增加稳定行
        all_lines.append(f"// ----- file: {rel} -----")
        all_lines.extend(body)
    return all_lines, len(files), len(all_lines)


def pick_front_back(lines: list[str]) -> list[str]:
    if len(lines) <= CHUNK_LINES * 2:
        return lines
    front = lines[:CHUNK_LINES]
    back = lines[-CHUNK_LINES:]
    # 前 30 页 + 后 30 页，合计 60 页（不再插入分隔行，避免多出一页）
    return front + back


def ensure_ends_with_closing(lines: list[str]) -> list[str]:
    """若最后一行不像结束，尽量从原逻辑已取文件末尾；再兜底补一个结束注释。"""
    if not lines:
        return ["// empty", "}"]
    last = lines[-1].rstrip()
    if last.endswith("}") or last.endswith("};") or last.endswith("*/"):
        return lines
    # 不伪造业务代码，仅追加明确的文档结束标记，避免以「{」收尾
    return lines + ["// END OF SOURCE DEPOSIT MATERIAL", "}"]


def add_code_line(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(11)
    run = p.add_run(text)
    set_run_font(run, size_pt=9)


def main() -> None:
    files = iter_source_files()
    all_lines, file_count, nonblank = collect_all_lines(files)
    deposit = ensure_ends_with_closing(pick_front_back(all_lines))

    doc = Document()
    setup_header_footer(doc)

    # 去掉 Normal 样式额外段后间距
    style = doc.styles["Normal"]
    style.font.size = Pt(9)
    style.font.name = "Courier New"

    # Document() 自带一个空段，复用为第一行，避免凭空多出 1 行/页
    if deposit:
        first = doc.paragraphs[0]
        pf = first.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(11)
        first.clear()
        run = first.add_run(deposit[0])
        set_run_font(run, size_pt=9)
        for line in deposit[1:]:
            add_code_line(doc, line)
    else:
        add_code_line(doc, "// empty")
        add_code_line(doc, "}")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)

    approx_pages = (len(deposit) + LINES_PER_PAGE - 1) // LINES_PER_PAGE
    print(f"源文件数: {file_count}")
    print(f"源程序量（建议填申请表）: {nonblank}行")
    print(f"交存文档行数: {len(deposit)} （约 {approx_pages} 页，按每页 {LINES_PER_PAGE} 行）")
    print(f"Wrote {OUT.relative_to(ROOT)}")
    print("注意: 申请表「源程序量」必须写成带单位，例如「{}行」。".format(nonblank))


if __name__ == "__main__":
    main()
