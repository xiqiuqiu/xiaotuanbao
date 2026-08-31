#!/usr/bin/env python3
"""将使用说明书 Markdown 转为软著打印用 docx（页眉 + 页码）。"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[3]
SRC = Path(__file__).resolve().parents[1] / "使用说明书.md"
OUT = Path(__file__).resolve().parents[1] / "output" / "小团宝旅行社运营管理系统V1.0-使用说明书.docx"

SOFTWARE_HEADER = "小团宝旅行社运营管理系统 V1.0"


def set_run_font(run, *, size_pt: float = 12, bold: bool = False, name: str = "宋体") -> None:
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:eastAsia"), name)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, size_pt=9)


def setup_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.clear()
    # 左软件名，右页码（用制表位）
    tab_stops = hp.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(16.5), WD_TAB_ALIGNMENT.RIGHT)
    run = hp.add_run(SOFTWARE_HEADER)
    set_run_font(run, size_pt=9)
    hp.add_run("\t")
    add_page_number(hp)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.clear()


def parse_md_lines(text: str) -> list[tuple[str, str]]:
    """粗粒度解析：返回 (kind, content)。kind: h1/h2/h3/p/table_row/blank/ul"""
    items: list[tuple[str, str]] = []
    lines = text.splitlines()
    i = 0
    # skip first H1 duplicate title handled specially
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            items.append(("blank", ""))
            i += 1
            continue
        if line.startswith("|") and "---" not in line:
            row_cells = [c.strip() for c in line.strip("|").split("|")]
            items.append(("table_row", " | ".join(row_cells)))
            i += 1
            continue
        if re.match(r"^\|?\s*-{3,}", line.replace("|", " ").strip()) or set(line.replace("|", "").strip()) <= {"-", ":", " "}:
            i += 1
            continue
        if line.startswith("# "):
            items.append(("h1", line[2:].strip()))
        elif line.startswith("## "):
            items.append(("h2", line[3:].strip()))
        elif line.startswith("### "):
            items.append(("h3", line[4:].strip()))
        elif line.startswith("#### "):
            items.append(("h3", line[5:].strip()))
        elif re.match(r"^\d+\.\s+", line.strip()) or line.strip().startswith(("- ", "* ")):
            items.append(("ul", re.sub(r"^[-*]\s+", "", line.strip())))
        elif line.strip() == "---":
            items.append(("blank", ""))
        else:
            items.append(("p", line.strip()))
        i += 1
    return items


def add_para(doc: Document, text: str, *, style: str = "Normal", size: float = 12, bold: bool = False, first_line: float | None = None) -> None:
    p = doc.add_paragraph()
    if style == "Title":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(3)
    pf.space_after = Pt(3)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if first_line is not None:
        pf.first_line_indent = Cm(first_line)
    run = p.add_run(text)
    set_run_font(run, size_pt=size, bold=bold, name="黑体" if bold and style != "Normal" else "宋体")
    if style == "Title":
        set_run_font(run, size_pt=size, bold=True, name="黑体")


def main() -> None:
    md = SRC.read_text(encoding="utf-8")
    doc = Document()
    setup_header_footer(doc)

    items = parse_md_lines(md)
    title_done = False
    for kind, content in items:
        if kind == "blank":
            continue
        if kind == "h1":
            if not title_done:
                add_para(doc, content, style="Title", size=16, bold=True)
                title_done = True
            else:
                add_para(doc, content, size=14, bold=True)
            continue
        if kind == "h2":
            add_para(doc, content, size=14, bold=True)
            continue
        if kind == "h3":
            add_para(doc, content, size=12, bold=True)
            continue
        if kind == "table_row":
            add_para(doc, content, size=10.5, first_line=0)
            continue
        if kind == "ul":
            add_para(doc, content, size=12, first_line=0.74)
            continue
        add_para(doc, content, size=12, first_line=0.74)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
